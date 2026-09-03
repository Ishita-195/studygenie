"""
StudyGenie Flask Server
- TF-IDF retrieval + Groq LLaMA 3 for answers, quiz, and summaries
- Persistent disk cache: restarts in <2 seconds regardless of PDF count
- Background auto-indexing: server available immediately on start
- Relative paths: works on any machine
"""

import os, sys, time, json, re, pickle, threading, hashlib
from pathlib import Path
import numpy as np   # ensure early import

# Fix Windows console encoding
if hasattr(sys.stdout, 'encoding') and sys.stdout.encoding != 'utf-8':
    sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', errors='replace', buffering=1)
    sys.stderr = open(sys.stderr.fileno(), mode='w', encoding='utf-8', errors='replace', buffering=1)

from flask import Flask, request, jsonify, Response
from flask_cors import CORS
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import fitz
from dotenv import load_dotenv

# ── Paths (all relative — works on any machine) ─────────────────────────────
BASE_DIR     = Path(__file__).parent.parent          # htdocs/ad_lab/
UPLOAD_DIR   = BASE_DIR / "uploads"
CACHE_DIR    = UPLOAD_DIR / ".cache"
ENV_FILE     = Path(__file__).parent / ".env"

load_dotenv(ENV_FILE)
UPLOAD_DIR.mkdir(exist_ok=True)
CACHE_DIR.mkdir(exist_ok=True)

# ── Config ───────────────────────────────────────────────────────────────────
GROQ_MODEL       = os.getenv("GROQ_MODEL",       "llama-3.3-70b-versatile")
GROQ_MODEL_FAST  = os.getenv("GROQ_MODEL_FAST",  "llama-3.1-8b-instant")   # smaller, separate quota
                                                 # (llama3-8b-8192 was decommissioned by Groq)
TOP_K       = 8      # more chunks → better context for complex questions
MIN_SCORE   = 0.03   # lower threshold → fewer "not found" on valid questions
CHUNK_WORDS = 300
CHUNK_STEP  = 200

# ── State ─────────────────────────────────────────────────────────────────────
per_file_data  = {}          # filename → {chunks, vectorizer, matrix, embeddings?}
per_file_sig   = {}          # filename → (mtime_ns, size)
index_status   = {}          # filename → "indexing" | "ready" | "error:<msg>"
_index_lock    = threading.Lock()

# ── Sentence-Transformer (semantic search) ────────────────────────────────────
_st_model       = None
_st_model_lock  = threading.Lock()
_st_model_ready = threading.Event()   # set when model is loaded

def _load_st_model():
    """Load sentence-transformer model in background thread."""
    global _st_model
    try:
        from sentence_transformers import SentenceTransformer
        print("[MODEL] Loading semantic model (all-MiniLM-L6-v2)…")
        m = SentenceTransformer('all-MiniLM-L6-v2')
        with _st_model_lock:
            _st_model = m
        _st_model_ready.set()
        print("[MODEL] ✅ Semantic search ready")
    except Exception as e:
        print(f"[MODEL] ⚠️ Could not load semantic model: {e} — falling back to TF-IDF")
        with _st_model_lock:
            _st_model = False          # False = unavailable (don't retry)
        _st_model_ready.set()

def _get_st_model():
    """Return loaded model (or None/False)."""
    with _st_model_lock:
        return _st_model

app = Flask(__name__)
CORS(app)

# ── Groq ─────────────────────────────────────────────────────────────────────
groq_client = None
_api_key = os.getenv("GROQ_API_KEY", "")
if _api_key and _api_key not in ("your_groq_api_key_here", ""):
    try:
        from groq import Groq
        groq_client = Groq(api_key=_api_key, timeout=45.0)
        print("[STARTUP] ✅ Groq client ready")
    except Exception as e:
        print(f"[STARTUP] ❌ Groq error: {e}")
else:
    print("[STARTUP] ⚠️  GROQ_API_KEY not set — quiz/answer will use fallback mode")
    print("[STARTUP]    Get a free key at https://console.groq.com/keys")
    print(f"[STARTUP]    Then add it to: {ENV_FILE}")


# ══════════════════════════════════════════════════════════════════════════════
#  CACHE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _file_sig(filepath: Path):
    try:
        st = filepath.stat()
        return (getattr(st, "st_mtime_ns", int(st.st_mtime * 1e9)), st.st_size)
    except OSError:
        return None

def _cache_path(filename: str) -> Path:
    safe = hashlib.md5(filename.encode()).hexdigest()
    return CACHE_DIR / f"{safe}.pkl"

def _save_cache(filename: str):
    try:
        entry = per_file_data[filename]
        payload = {
            "chunks":     entry["chunks"],
            "vectorizer": entry["vectorizer"],
            "matrix":     entry["matrix"],
            "sig":        per_file_sig[filename],
            "filename":   filename,
        }
        if "embeddings" in entry:
            payload["embeddings"] = entry["embeddings"]
        with open(_cache_path(filename), "wb") as f:
            pickle.dump(payload, f, protocol=4)
    except Exception as e:
        print(f"[CACHE] Save failed for {filename}: {e}")

def _load_cache(filename: str, filepath: Path) -> bool:
    cp = _cache_path(filename)
    if not cp.exists():
        return False
    try:
        current_sig = _file_sig(filepath)
        with open(cp, "rb") as f:
            cached = pickle.load(f)
        if cached.get("sig") != current_sig:
            cp.unlink(missing_ok=True)   # stale cache
            return False
        entry = {
            "chunks":     cached["chunks"],
            "vectorizer": cached["vectorizer"],
            "matrix":     cached["matrix"],
        }
        if "embeddings" in cached:
            entry["embeddings"] = cached["embeddings"]
        per_file_data[filename] = entry
        per_file_sig[filename] = current_sig
        has_sem = "embeddings" in cached
        print(f"[CACHE] ✅ {filename} ({len(cached['chunks'])} chunks, semantic={'yes' if has_sem else 'no'})")
        return True
    except Exception as e:
        print(f"[CACHE] Load failed for {filename}: {e}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT EXTRACTION  (PDF + OCR fallback + DOCX)
# ══════════════════════════════════════════════════════════════════════════════

# Detect OCR availability once at startup
_OCR_AVAILABLE = False
try:
    import pytesseract
    from PIL import Image
    # Common Windows install path for the Tesseract engine
    _tess_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for _p in _tess_paths:
        if os.path.exists(_p):
            pytesseract.pytesseract.tesseract_cmd = _p
            break
    pytesseract.get_tesseract_version()   # raises if engine missing
    _OCR_AVAILABLE = True
    print("[STARTUP] ✅ OCR available (scanned PDFs supported)")
except Exception:
    print("[STARTUP] ⚠️ OCR not available — scanned/image PDFs will be skipped")
    print("[STARTUP]    Install Tesseract: https://github.com/UB-Mannheim/tesseract/wiki")

try:
    import docx as _docx          # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _ocr_pdf(doc) -> str:
    """Run OCR on every page of a PDF (used when normal extraction finds no text)."""
    if not _OCR_AVAILABLE:
        return ""
    text_parts = []
    for page in doc:
        pix = page.get_pixmap(dpi=200)            # render page to image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        text_parts.append(pytesseract.image_to_string(img))
    return " ".join(text_parts)


def _extract_text(filepath: Path):
    """
    Extract text from a document. Supports:
      - .pdf  (text-based → OCR fallback for scanned PDFs)
      - .docx (Microsoft Word)
    Returns (text, page_count).
    """
    ext = filepath.suffix.lower()

    if ext == ".docx":
        if not _DOCX_AVAILABLE:
            raise ValueError("DOCX support not installed (pip install python-docx).")
        d = _docx.Document(str(filepath))
        paras  = [p.text for p in d.paragraphs if p.text.strip()]
        # Include table cell text too
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paras.append(cell.text)
        return " ".join(paras), 1

    # Default: PDF
    doc = fitz.open(str(filepath))
    try:
        pages    = len(doc)
        all_text = "".join(page.get_text() + " " for page in doc)
        # If almost no text was extracted, the PDF is likely scanned → OCR
        if len(all_text.split()) < 20 and _OCR_AVAILABLE:
            print(f"[INDEX] Low text — running OCR on {pages} pages…")
            all_text = _ocr_pdf(doc)
        return all_text, pages
    finally:
        doc.close()


_TOPIC_STOP = {
    "edition","fifth","fourth","third","second","first","chapter","section","figure",
    "table","contents","index","preface","introduction","summary","references","among",
    "these","those","there","their","which","while","about","using","based","given",
    "however","therefore","university","press","copyright","reserved","rights","author",
    "professor","january","february","march","april","august","september","october",
    "november","december","monday","other","number","example","following","above","below",
    "basic","bibliography","glossary","appendix","volume","part","unit","page","note","notes",
    # common capitalised sentence-starters / function words
    "that","this","with","when","from","each","they","then","there","what","where","while",
    "also","such","here","your","have","will","would","could","should","into","more","than",
    "some","most","many","both","only","even","still","much","very","well","like","just",
    "make","made","used","over","after","before","because","between","through","since","during",
    "every","then","they","them","then","being","does","done","each","then","when","then",
    "however","thus","therefore","then","then","once","upon","when","whenever","whereas"
}

def _important_terms(filename: str, top_n: int = 20) -> list:
    """
    Return the document's most important terms, scored by TF-IDF weight.
    These are the terms that genuinely characterise THIS document
    (not just frequent words) — used to anchor the concept map.
    """
    data = per_file_data.get(filename)
    if not data:
        return []
    try:
        vec   = data["vectorizer"]
        mat   = data["matrix"]
        names = vec.get_feature_names_out()
        scores = mat.sum(axis=0).A1                      # total TF-IDF per term
        ranked = sorted(zip(scores, names), key=lambda x: -x[0])
        out, seen = [], set()
        for _, term in ranked:
            t = term.strip()
            # keep multi-word phrases and meaningful single words; drop junk
            words = t.split()
            if any(w in _TOPIC_STOP for w in words):
                continue
            if len(t) < 4 or t in seen or t.isdigit():
                continue
            seen.add(t)
            out.append(t.title())
            if len(out) >= top_n:
                break
        return out
    except Exception as e:
        print(f"[IMPORTANT-TERMS] {e}")
        return []


def _extract_topic_words(chunks: list, limit: int) -> list:
    """
    Extract meaningful topic terms from chunks, skipping junk.
    Key heuristic: only keep terms that RECUR in the document — real concepts
    repeat many times, whereas author names / cover-page words appear once.
    """
    full = " ".join(chunks)
    candidates, seen = [], set()
    for m in re.findall(r'\b[A-Z][a-zA-Z]{3,}(?:\s+[A-Z][a-zA-Z]{3,}){0,2}', full):
        first = m.split()[0].lower()
        key = m.lower()
        if first in _TOPIC_STOP or key in seen:
            continue
        seen.add(key)
        # Count case-insensitive occurrences of this term across the document
        freq = len(re.findall(r'\b' + re.escape(m) + r'\b', full, re.IGNORECASE))
        candidates.append((freq, m.strip()))

    # Prefer the most frequently-recurring terms (real concepts), require freq >= 2
    recurring = [t for f, t in sorted(candidates, key=lambda x: -x[0]) if f >= 2]
    result = recurring[:limit]
    if len(result) < limit:   # top up with single-occurrence terms if needed
        result += [t for f, t in candidates if t not in result][:limit - len(result)]
    return result[:limit]


SUPPORTED_EXTS = (".pdf", ".docx")

def _list_documents():
    """All supported documents in the upload folder (PDF + DOCX)."""
    docs = []
    for ext in SUPPORTED_EXTS:
        docs.extend(UPLOAD_DIR.glob(f"*{ext}"))
    return docs


# ══════════════════════════════════════════════════════════════════════════════
#  INDEXING
# ══════════════════════════════════════════════════════════════════════════════

def _index_file(filename: str, filepath: Path):
    """Extract text, chunk, vectorize, cache. Returns (chunk_count, elapsed)."""
    start = time.time()
    size_mb = filepath.stat().st_size / 1_048_576
    print(f"\n[INDEX] {filename} ({size_mb:.1f} MB)")

    all_text, pages = _extract_text(filepath)

    words = all_text.split()
    print(f"[INDEX] {pages} pages, {len(words)} words")

    if len(words) < 20:
        hint = "" if _OCR_AVAILABLE else " (install Tesseract OCR to read scanned PDFs)"
        raise ValueError(f"Too little text ({len(words)} words). Document may be scanned/image-based{hint}.")

    chunks = [
        " ".join(words[i:i + CHUNK_WORDS]).strip()
        for i in range(0, len(words), CHUNK_STEP)
        if " ".join(words[i:i + CHUNK_WORDS]).strip()
    ]

    if not chunks:
        raise ValueError("No text chunks could be created from this PDF.")

    vectorizer   = TfidfVectorizer(max_features=8000, stop_words="english", ngram_range=(1, 2))
    tfidf_matrix = vectorizer.fit_transform(chunks)

    # Generate semantic embeddings (if model ready; skip if still loading)
    embeddings = None
    st = _get_st_model()
    if st and st is not False:
        try:
            t_emb = time.time()
            embeddings = st.encode(chunks, batch_size=64,
                                   show_progress_bar=False, convert_to_numpy=True)
            print(f"[INDEX] Semantic embeddings: {embeddings.shape} in {time.time()-t_emb:.1f}s")
        except Exception as e:
            print(f"[INDEX] Embedding error: {e}")

    entry = {"chunks": chunks, "vectorizer": vectorizer, "matrix": tfidf_matrix}
    if embeddings is not None:
        entry["embeddings"] = embeddings

    with _index_lock:
        per_file_data[filename] = entry
        per_file_sig[filename]  = _file_sig(filepath)
        index_status[filename]  = "ready"

    _save_cache(filename)
    elapsed = time.time() - start
    print(f"[INDEX] ✅ {filename} — {len(chunks)} chunks in {elapsed:.2f}s")
    return len(chunks), elapsed


def _ensure_indexed(filename: str):
    """
    Make sure a file is indexed. Returns (True, None) or (False, error_str).
    Order: memory → disk cache → re-index from PDF.
    """
    if not filename:
        return False, "No filename provided."

    filename = os.path.basename(filename).strip()
    filepath = UPLOAD_DIR / filename

    # Already in memory and file unchanged?
    with _index_lock:
        if filename in per_file_data:
            if _file_sig(filepath) == per_file_sig.get(filename):
                return True, None
            # File changed — drop stale index
            del per_file_data[filename]

    if not filepath.exists():
        available = [f.name for f in _list_documents()]
        return False, f"File '{filename}' not found on server. Available: {available}"

    # Try loading from disk cache first (fast)
    if _load_cache(filename, filepath):
        with _index_lock:
            index_status[filename] = "ready"
        return True, None

    # Full re-index
    with _index_lock:
        index_status[filename] = "indexing"
    try:
        _index_file(filename, filepath)
        return True, None
    except Exception as e:
        with _index_lock:
            index_status[filename] = f"error:{e}"
        return False, f"Indexing failed for '{filename}': {e}"


def _auto_index_all_bg():
    """
    Run in background:
    1. Load all PDFs from disk cache (fast)
    2. Wait for semantic model to be ready
    3. Generate embeddings for any doc missing them
    """
    pdfs = _list_documents()
    print(f"[STARTUP] Background indexing {len(pdfs)} documents…")
    for pdf in pdfs:
        fn = pdf.name
        with _index_lock:
            if fn in per_file_data:
                continue
            index_status[fn] = "indexing"
        try:
            if not _load_cache(fn, pdf):
                _index_file(fn, pdf)
            with _index_lock:
                index_status[fn] = "ready"
        except Exception as e:
            with _index_lock:
                index_status[fn] = f"error:{e}"
            print(f"[STARTUP] ❌ {fn}: {e}")

    ready = sum(1 for v in index_status.values() if v == "ready")
    print(f"[STARTUP] ✅ Indexing complete — {ready}/{len(pdfs)} ready")

    # Wait for semantic model, then backfill missing embeddings
    print("[STARTUP] Waiting for semantic model to backfill embeddings…")
    _st_model_ready.wait(timeout=120)
    st = _get_st_model()
    if not st or st is False:
        print("[STARTUP] Semantic model unavailable — skipping embedding backfill")
        return

    filled = 0
    with _index_lock:
        filenames_snapshot = list(per_file_data.keys())

    for fn in filenames_snapshot:
        if "embeddings" in per_file_data.get(fn, {}):
            continue                    # already has embeddings
        chunks = per_file_data.get(fn, {}).get("chunks", [])
        if not chunks:
            continue
        try:
            emb = st.encode(chunks, batch_size=64,
                            show_progress_bar=False, convert_to_numpy=True)
            with _index_lock:
                if fn in per_file_data:
                    per_file_data[fn]["embeddings"] = emb
            _save_cache(fn)
            filled += 1
            print(f"[STARTUP] Embeddings backfilled: {fn}")
        except Exception as e:
            print(f"[STARTUP] Embedding error for {fn}: {e}")

    print(f"[STARTUP] ✅ Semantic embeddings ready — {filled} docs updated")


# ══════════════════════════════════════════════════════════════════════════════
#  ANSWER / QUIZ HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _retrieve_chunks(filename: str, query: str, k=TOP_K):
    """TF-IDF retrieval (keyword-based). Used by /ask for backward compat."""
    data        = per_file_data[filename]
    q_vec       = data["vectorizer"].transform([query])
    scores      = cosine_similarity(q_vec, data["matrix"]).flatten()
    top_idx     = scores.argsort()[::-1][:k]
    top_score   = float(scores[top_idx[0]]) if len(top_idx) else 0
    relevant    = [data["chunks"][i] for i in top_idx if scores[i] >= MIN_SCORE]
    return relevant, top_score


def _retrieve_smart(filename: str, query: str, k=TOP_K):
    """
    Smart retrieval for /chat:
    1. Small docs (≤5 chunks) → return ALL chunks, no filtering
    2. Semantic search if embeddings available (sentence-transformers)
    3. TF-IDF fallback, but NEVER return empty for a valid doc
    Always returns (chunks, top_score) — never empty if doc has content.
    """
    data   = per_file_data[filename]
    chunks = data["chunks"]

    # Rule 1: tiny document → use everything, let Groq decide
    if len(chunks) <= 5:
        return chunks, 1.0

    # Rule 2: try semantic search
    st = _get_st_model()
    if st and st is not False:
        if "embeddings" not in data:
            # Generate embeddings on-the-fly (doc was cached without them)
            try:
                emb = st.encode(chunks, batch_size=64,
                                show_progress_bar=False, convert_to_numpy=True)
                with _index_lock:
                    data["embeddings"] = emb
                _save_cache(filename)
                print(f"[RETRIEVAL] Generated missing embeddings for {filename}")
            except Exception as e:
                print(f"[RETRIEVAL] Embedding generation failed: {e}")

        if "embeddings" in data:
            q_emb  = st.encode([query], convert_to_numpy=True)
            scores = cosine_similarity(q_emb, data["embeddings"]).flatten()
            top_i  = scores.argsort()[::-1][:k]
            top_s  = float(scores[top_i[0]])
            # Return top-k; no hard cutoff (semantic scores are generally lower)
            relevant = [chunks[i] for i in top_i]
            return relevant, top_s

    # Rule 3: TF-IDF fallback — but always return top-k regardless of score
    q_vec  = data["vectorizer"].transform([query])
    scores = cosine_similarity(q_vec, data["matrix"]).flatten()
    top_i  = scores.argsort()[::-1][:k]
    top_s  = float(scores[top_i[0]]) if len(top_i) else 0
    # No MIN_SCORE filter — return top-k regardless
    relevant = [chunks[i] for i in top_i]
    return relevant, top_s


def _extractive_answer(question: str, chunks: list) -> str:
    if not chunks:
        return "This topic is not found in the uploaded document."
    kw = {w for w in re.findall(r"\b[a-zA-Z]{4,}\b", question.lower())} - {
        "what","when","where","which","there","their","they","would","could","should",
        "about","from","this","that","with","have","been","into","used","using","make",
        "does","define","explain","describe","please","give","list","name","tell"
    }
    scored = []
    for chunk in chunks:
        for sent in re.split(r"(?<=[.!?])\s+", chunk):
            sent = sent.strip()
            if len(sent) < 25: continue
            hits = sum(1 for k in kw if k in sent.lower())
            scored.append((hits, len(sent), sent))
    if not scored:
        return chunks[0][:600]
    scored.sort(key=lambda x: (x[0], x[1]), reverse=True)
    best = scored[0][2]
    extras = [s[2] for s in scored[1:4] if s[2] != best][:2]
    result = f"Answer: {best}"
    if extras:
        result += "\nKey points:\n" + "\n".join(f"- {e}" for e in extras)
    return result


def _parse_json_response(text: str):
    """Robustly extract a JSON array or object from LLM output."""
    for fn in [
        lambda t: json.loads(t.strip()),
        lambda t: json.loads(re.search(r'```(?:json)?\s*([\s\S]*?)```', t).group(1)),
        lambda t: json.loads(re.search(r'(\[[\s\S]*\])', t).group(1)),
        lambda t: json.loads(re.search(r'(\{[\s\S]*\})', t).group(1)),
    ]:
        try:
            result = fn(text)
            if result: return result
        except:
            pass
    return None


def _make_questions_from_text(text: str, n: int, distractor_pool: list) -> list:
    """
    Core question factory: given a block of text, extract fill-in-the-blank
    questions. Works on ANY amount of text, even a single sentence.
    """
    import random

    STOP = {
        "this","that","with","from","have","been","they","their","will","were","also",
        "more","than","into","over","such","each","when","which","there","about","some",
        "these","those","other","while","after","before","through","between","because",
        "however","therefore","thus","hence","although","since","during","page","figure",
        "called","known","used","using","made","make","based","given","found","shown",
        "first","second","third","chapter","section","equation","table","note","example"
    }

    sentences = re.split(r'(?<=[.!?])\s+', text)
    questions = []
    used_answers = set()

    for sent in sentences:
        if len(questions) >= n:
            break
        sent = sent.strip()
        if len(sent) < 30:
            continue

        words = sent.split()
        # Pick keyword: prefer longer words not in stop list
        candidates = [
            w for w in words
            if len(re.sub(r'[^a-zA-Z]', '', w)) >= 4
            and re.sub(r'[^a-zA-Z]', '', w).lower() not in STOP
            and re.sub(r'[^a-zA-Z]', '', w) not in used_answers
            and re.sub(r'[^a-zA-Z]', '', w).isalpha()
        ]
        if not candidates:
            continue

        answer_word  = max(candidates, key=lambda w: len(re.sub(r'[^a-zA-Z]', '', w)))
        answer_clean = re.sub(r'[^a-zA-Z]', '', answer_word)
        if len(answer_clean) < 4:
            continue

        q_text = sent.replace(answer_word, "______", 1)

        # Build distractors from pool, similar length
        pool = [t for t in distractor_pool
                if t != answer_clean and abs(len(t) - len(answer_clean)) <= 5]
        if len(pool) < 3:
            pool = [t for t in distractor_pool if t != answer_clean]
        if len(pool) < 3:
            pool = ["protocol", "network", "system", "algorithm",
                    "device", "function", "interface", "method",
                    "process", "signal", "layer", "packet"]
        distractors = random.sample(pool[:100], min(3, len(pool)))
        opts = [answer_clean] + distractors
        random.shuffle(opts)

        questions.append({
            "question": q_text,
            "options":  opts[:4],
            "answer":   opts.index(answer_clean)
        })
        used_answers.add(answer_clean)

    return questions


def _fallback_quiz(chunks: list, n=5) -> list:
    """
    Guaranteed to return exactly n questions from ANY document, no matter how small.
    Strategy:
      1. Try definition-style sentences from content chunks (skipping front matter)
      2. If still short, use ALL sentences from the full document
      3. If still short, repeat/rephrase with different blanked words
      4. Last resort: pad with generic document-topic questions
    """
    import random

    # Build distractor pool from the whole document
    STOP = {
        "this","that","with","from","have","been","they","their","will","were","also",
        "more","than","into","over","such","each","when","which","there","about","some",
        "these","those","other","while","after","before","through","between","because",
        "however","therefore","thus","hence","although","since","during","page","figure",
        "called","known","used","using","made","make","based","given","found","shown",
        "first","second","third","chapter","section","equation","table","note","example"
    }
    pool = list(dict.fromkeys(
        re.sub(r'[^a-zA-Z]', '', w)
        for ch in chunks
        for w in ch.split()
        if len(re.sub(r'[^a-zA-Z]', '', w)) >= 4
        and re.sub(r'[^a-zA-Z]', '', w).isalpha()
        and re.sub(r'[^a-zA-Z]', '', w).lower() not in STOP
    ))[:300]

    # Use only content-rich chunks (skips TOC / index / cover) for the fallback too
    clean      = _quiz_chunks(chunks, k=12) or chunks
    full_text  = " ".join(clean)

    # Pass 1: definition-style questions from cleaned content
    questions = _make_questions_from_text(full_text, n, pool)

    # Pass 2: if not enough, use the entire document text
    if len(questions) < n:
        extra = _make_questions_from_text(full_text, n - len(questions), pool)
        # Avoid duplicates
        existing = {q["question"] for q in questions}
        questions += [q for q in extra if q["question"] not in existing]

    # Pass 3: if STILL short (very tiny doc), create questions by cycling through
    # all words in the document as blanks
    if len(questions) < n:
        words_in_doc = [
            re.sub(r'[^a-zA-Z]', '', w) for w in full_text.split()
            if len(re.sub(r'[^a-zA-Z]', '', w)) >= 5
            and re.sub(r'[^a-zA-Z]', '', w).isalpha()
            and re.sub(r'[^a-zA-Z]', '', w).lower() not in STOP
        ]
        used_so_far = {q["options"][q["answer"]] for q in questions}
        for ans in words_in_doc:
            if len(questions) >= n:
                break
            if ans in used_so_far:
                continue
            distractors = random.sample(
                [w for w in pool if w != ans][:80] or ["network","system","method","process"],
                min(3, len([w for w in pool if w != ans]) or 3)
            )
            opts = [ans] + distractors
            random.shuffle(opts)
            # Build a context snippet
            idx = full_text.find(ans)
            snippet = full_text[max(0, idx-60):idx+len(ans)+60].strip()
            snippet = re.sub(r'\s+', ' ', snippet)
            question_text = snippet.replace(ans, "______", 1) if ans in snippet else f"The document mentions ______."
            questions.append({
                "question": question_text,
                "options":  opts[:4],
                "answer":   opts.index(ans)
            })
            used_so_far.add(ans)

    # Pass 4: absolute last resort — pad with generic Qs if document has almost no text
    GENERIC = [
        {"question": "What is the primary subject of this document?",
         "options":  ["The topic described in this document","Historical events","Mathematical proofs","Programming code"],
         "answer":   0},
        {"question": "What type of content does this document contain?",
         "options":  ["Technical or academic content","Fiction story","Legal contract","News article"],
         "answer":   0},
        {"question": "What would you use this document for?",
         "options":  ["Learning and study","Entertainment","Legal reference","Financial planning"],
         "answer":   0},
        {"question": "Who is the likely audience for this document?",
         "options":  ["Students or professionals in the field","Children","General public only","Musicians"],
         "answer":   0},
        {"question": "What is the best way to use this document?",
         "options":  ["Read carefully and take notes","Skim for entertainment","Use as a recipe guide","Print and discard"],
         "answer":   0},
    ]
    idx = 0
    while len(questions) < n and idx < len(GENERIC):
        questions.append(GENERIC[idx])
        idx += 1

    # Ensure every fallback question carries difficulty + explanation fields
    for q in questions:
        q.setdefault("difficulty", "medium")
        q.setdefault("explanation", "")

    return questions[:n]


# ══════════════════════════════════════════════════════════════════════════════
#  ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/upload-and-index", methods=["POST"])
def upload_and_index():
    data     = request.get_json(silent=True) or {}
    filename = os.path.basename(data.get("filename", "")).strip()
    if not filename:
        return jsonify({"error": "No filename provided"}), 400
    ok, err = _ensure_indexed(filename)
    if not ok:
        return jsonify({"error": err}), 404
    chunks = len(per_file_data[filename]["chunks"])
    return jsonify({"status": "indexed", "filename": filename, "chunks": chunks,
                    "time_seconds": 0})


@app.route("/chat", methods=["POST"])
def chat():
    """
    Conversational endpoint with:
    - Full conversation history for follow-up questions
    - Document-grounded answers (no hallucination)
    - Structured, human-readable responses
    - Specific question answering (not document dumps)
    """
    data      = request.get_json(silent=True) or {}
    question  = data.get("question", "").strip()
    filename  = os.path.basename(data.get("filename", "")).strip()
    history   = data.get("history", [])      # [{role, content}, ...]
    print(f"\n[CHAT] '{question[:70]}' | file='{filename}' | history={len(history)}")

    if not question:
        return jsonify({"status": "error", "answer": "Please enter a question."})
    if not filename:
        return jsonify({"status": "error", "answer": "No document selected."})

    ok, err = _ensure_indexed(filename)
    if not ok:
        return jsonify({"status": "error", "answer": err})

    # Smart retrieval — semantic first, TF-IDF fallback, NEVER returns empty for valid doc
    relevant, top_score = _retrieve_smart(filename, question)
    confidence = min(95, int(top_score * 180))

    # Extractive fallback (no Groq needed)
    extractive = _extractive_answer(question, relevant) if relevant else None

    if not relevant:
        return jsonify({
            "status":     "ok",
            "answer":     "⚠️ This document doesn't appear to contain any extractable text. It may be a scanned PDF or image-only file. Please try uploading a text-based PDF.",
            "confidence": 0
        })

    context = "\n\n---\n\n".join(relevant)

    if groq_client:
        try:
            # Comprehensive system prompt for structured, accurate answers
            system = f"""You are StudyGenie, an intelligent document AI assistant. Your job is to help users understand their uploaded document by answering questions accurately and clearly.

DOCUMENT CONTENT (extracted from uploaded PDF):
===
{context}
===

YOUR RULES:
1. Answer the user's SPECIFIC question using ONLY the document content above
2. NEVER copy-paste raw text from the document — always rephrase, explain, and structure
3. NEVER invent facts not present in the document
4. Format EVERY response using this structure (skip sections that don't apply):

**Answer**
[Direct, clear answer to the question in 1-3 sentences]

**Explanation**
[Detailed explanation with context from the document]

**Key Points**
- [Important highlight 1]
- [Important highlight 2]
- [Important highlight 3]

**Source**
[Brief note on which part of the document this comes from]

5. If the document has very limited text, work with what's available and note the limitation
6. If asked to "list", "summarize", or "explain" — always produce a proper structured list/summary
7. Use **bold** for technical terms
8. Be conversational, clear, and easy to understand even for complex topics
9. For follow-up questions, use the conversation history for context"""

            # Build messages with conversation history
            messages = [{"role": "system", "content": system}]
            for h in history[-12:]:
                if h.get("role") in ("user", "assistant") and h.get("content"):
                    messages.append({"role": h["role"], "content": h["content"]})
            messages.append({"role": "user", "content": question})

            resp   = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=messages,
                temperature=0.25,
                max_tokens=900
            )
            answer = resp.choices[0].message.content.strip()
            print(f"[CHAT] ✅ Groq answered ({len(answer)} chars)")
        except Exception as e:
            print(f"[CHAT] Groq error: {e}")
            answer = extractive or "Unable to generate answer."
    else:
        answer = extractive or "AI model not configured. Please set GROQ_API_KEY."

    return jsonify({
        "status":     "ok",
        "answer":     answer,
        "confidence": confidence,
        "sources":    [filename]
    })


@app.route("/ask", methods=["POST"])
def ask():
    data     = request.get_json(silent=True) or {}
    question = data.get("question", "").strip()
    filename = os.path.basename(data.get("filename", "")).strip()
    print(f"\n[ASK] '{question[:60]}' | file='{filename}'")

    if not filename:
        return jsonify({"status":"error","answer":"No filename provided.","sources":[],"confidence":0})

    ok, err = _ensure_indexed(filename)
    if not ok:
        return jsonify({"status":"error","answer":err,"sources":[],"confidence":0})

    relevant, top_score = _retrieve_chunks(filename, question)

    # Second-pass: if nothing found above MIN_SCORE, widen the net to top 5 chunks regardless of score
    if not relevant:
        data2     = per_file_data[filename]
        q_vec     = data2["vectorizer"].transform([question])
        scores    = cosine_similarity(q_vec, data2["matrix"]).flatten()
        top_idx   = scores.argsort()[::-1][:5]
        top_score = float(scores[top_idx[0]]) if len(top_idx) else 0
        if top_score > 0.01:                        # at least loosely related
            relevant = [data2["chunks"][i] for i in top_idx]

    confidence = min(95, int(top_score * 220))

    if not relevant:
        return jsonify({"status":"ok","answer":"This topic is not found in the uploaded document.",
                        "sources":[filename],"confidence":0})

    context    = "\n---\n".join(relevant)
    extractive = _extractive_answer(question, relevant)

    if groq_client:
        try:
            resp = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role":"system","content":(
                        "You are StudyGenie, a helpful academic assistant. "
                        "Answer the user's question using ONLY the provided document context. "
                        "Give a thorough, detailed answer — not just a one-liner. "
                        "Explain concepts clearly as if teaching a student. "
                        "If the answer is not in the context at all, say: "
                        "'This topic is not covered in the uploaded document.' "
                        "Format your answer as:\n"
                        "Answer: <clear explanation in 2-4 sentences>\n"
                        "Key points:\n- <important detail>\n- <important detail>\n- <important detail>"
                    )},
                    {"role":"user","content":f"Context:\n{context}\n\nQuestion: {question}"}
                ],
                temperature=0.15, max_tokens=700
            )
            answer = resp.choices[0].message.content
        except Exception as e:
            print(f"[ASK] Groq error: {e}")
            answer = extractive
    else:
        answer = extractive

    return jsonify({"status":"ok","answer":answer,"sources":[filename],"confidence":confidence})


@app.route("/ask-stream", methods=["POST"])
def ask_stream():
    data     = request.get_json(silent=True) or {}
    question = data.get("question","").strip()
    filename = os.path.basename(data.get("filename","")).strip()

    def err_stream(msg):
        yield f'data: {json.dumps({"token":msg})}\n\n'
        yield 'data: [DONE]\n\n'

    if not filename:
        return Response(err_stream("No filename provided."), mimetype="text/event-stream")

    ok, err_msg = _ensure_indexed(filename)
    if not ok:
        return Response(err_stream(err_msg), mimetype="text/event-stream")

    relevant, _ = _retrieve_chunks(filename, question)
    if not relevant:
        return Response(err_stream("This topic is not found in the uploaded document."),
                        mimetype="text/event-stream")

    context    = "\n---\n".join(relevant)
    extractive = _extractive_answer(question, relevant)

    def generate():
        if groq_client:
            try:
                stream = groq_client.chat.completions.create(
                    model=GROQ_MODEL,
                    messages=[
                        {"role":"system","content":"You are StudyGenie. Answer from context only."},
                        {"role":"user","content":f"Context:\n{context}\n\nQuestion: {question}"}
                    ],
                    stream=True, max_tokens=500
                )
                for chunk in stream:
                    token = chunk.choices[0].delta.content
                    if token:
                        yield f'data: {json.dumps({"token":token})}\n\n'
            except Exception as e:
                for w in extractive.split():
                    yield f'data: {json.dumps({"token":w+" "})}\n\n'
        else:
            for w in extractive.split():
                yield f'data: {json.dumps({"token":w+" "})}\n\n'
        yield 'data: [DONE]\n\n'

    return Response(generate(), mimetype="text/event-stream")


def _is_garbage_chunk(text: str) -> bool:
    """Detect table-of-contents / index / cover / reference-list chunks
    that make terrible quiz material."""
    tokens = text.split()
    if len(tokens) < 25:
        return True
    # TOC/index pages are dense with page numbers
    digit_tokens = sum(1 for t in tokens if any(c.isdigit() for c in t))
    if digit_tokens / len(tokens) > 0.22:
        return True
    upper_head = text[:250].upper()
    if any(k in upper_head for k in ("CONTENTS", "TABLE OF CONTENTS", "INDEX", "REFERENCES", "BIBLIOGRAPHY")):
        return True
    if text.count("....") > 2 or text.count(". . .") > 2:   # dot leaders
        return True
    # Needs real prose: at least a few sentence-ending periods
    if text.count(". ") < 3:
        return True
    return False


def _quiz_chunks(chunks: list, k: int = 8) -> list:
    """Select content-rich chunks spread evenly across the whole document,
    skipping front-matter and TOC/index garbage — for balanced coverage."""
    # Skip obvious front matter (cover/title/copyright)
    body = chunks[min(3, len(chunks) - 1):] if len(chunks) > 6 else chunks
    good = [c for c in body if not _is_garbage_chunk(c)]
    if len(good) < 2:
        good = [c for c in chunks if not _is_garbage_chunk(c)] or chunks
    if len(good) <= k:
        return good
    # Spread selection evenly across the document for balanced coverage
    step = max(1, len(good) // k)
    return good[::step][:k]


@app.route("/generate-quiz", methods=["POST"])
def generate_quiz():
    data          = request.get_json(silent=True) or {}
    filename      = os.path.basename(data.get("filename","")).strip()
    num_questions = max(1, min(int(data.get("num_questions", 5)), 15))
    print(f"\n[QUIZ] {num_questions}q from '{filename}'")

    ok, err = _ensure_indexed(filename)
    if not ok:
        return jsonify({"error": err, "questions": []})

    chunks  = per_file_data[filename]["chunks"]
    # Select content-rich chunks spread across the document (skips TOC/index/cover)
    sampled = _quiz_chunks(chunks, k=8)
    context = "\n\n---\n\n".join(sampled)

    # Cache (v2 key busts old low-quality cached quizzes)
    quiz_cache_path = CACHE_DIR / (hashlib.md5(filename.encode()).hexdigest() + f"_quizv2_{num_questions}.json")
    if quiz_cache_path.exists():
        try:
            with open(quiz_cache_path, encoding="utf-8") as f:
                cached_quiz = json.load(f)
            if cached_quiz.get("questions"):
                print(f"[QUIZ] ✅ Loaded {len(cached_quiz['questions'])} questions from cache")
                return jsonify(cached_quiz)
        except:
            pass

    if not groq_client:
        fb = _fallback_quiz(chunks, num_questions)
        return jsonify({"status":"ok","questions":fb,"mode":"fallback"})

    # Concept-focused, education-grade prompt with explanations + difficulty mix
    prompt = (
        f"You are an expert exam author. Create exactly {num_questions} high-quality "
        "multiple-choice questions that test genuine understanding of the KEY CONCEPTS "
        "in the document content below.\n\n"
        "STRICT REQUIREMENTS:\n"
        "1. Test comprehension, application, and analysis — NOT trivial recall or fill-in-the-blank\n"
        "2. Use a MIX of question types: conceptual, definition, application, scenario-based, "
        "cause-and-effect, and comparison (where the content allows)\n"
        "3. Include a MIX of difficulty levels: some easy, some medium, some hard\n"
        "4. Each question covers a DIFFERENT topic/section — never repeat the same point\n"
        "5. Each question has exactly 4 options. ALL distractors must be plausible and "
        "believable (no obviously wrong joke options). Only ONE option is correct.\n"
        "6. Questions must be grammatically correct, clear, and unambiguous\n"
        "7. NEVER copy a sentence verbatim and NEVER reference 'the text/document/passage/page'\n"
        "8. For each question include a 1-2 sentence explanation of WHY the answer is correct\n\n"
        "Return ONLY a valid JSON array (no extra text):\n"
        '[{"question":"...","options":["opt1","opt2","opt3","opt4"],"answer":0,'
        '"difficulty":"easy","explanation":"..."}]\n\n'
        f"Document content:\n{context}"
    )

    # Quality-first: try the large model first, fall back to the fast model on rate-limit
    for model in [GROQ_MODEL, GROQ_MODEL_FAST]:
        try:
            resp   = groq_client.chat.completions.create(
                model=model, messages=[{"role":"user","content":prompt}],
                temperature=0.6, max_tokens=3000
            )
            raw    = resp.choices[0].message.content.strip()
            parsed = _parse_json_response(raw)
            if parsed and isinstance(parsed, list):
                valid = []
                for q in parsed:
                    if (isinstance(q, dict)
                            and isinstance(q.get("question"), str)
                            and isinstance(q.get("options"), list)
                            and len(q.get("options")) == 4
                            and isinstance(q.get("answer"), int)
                            and 0 <= q.get("answer") <= 3):
                        # Normalise optional fields
                        q["difficulty"]  = str(q.get("difficulty", "medium")).lower()
                        if q["difficulty"] not in ("easy", "medium", "hard"):
                            q["difficulty"] = "medium"
                        q["explanation"] = str(q.get("explanation", "")).strip()
                        valid.append(q)
                if valid:
                    result = {"status":"ok","questions":valid[:num_questions],"model":model}
                    print(f"[QUIZ] ✅ {len(valid)} concept questions via {model}")
                    try:
                        with open(quiz_cache_path, "w", encoding="utf-8") as f:
                            json.dump(result, f)
                    except: pass
                    return jsonify(result)
            print(f"[QUIZ] ⚠️ Parse/validate failed with {model}, trying next")
        except Exception as e:
            print(f"[QUIZ] {model} error: {e}")
            continue   # any error → try the next model before giving up to fallback

    fb = _fallback_quiz(chunks, num_questions)
    return jsonify({"status":"ok","questions":fb,"mode":"fallback"})


@app.route("/summarize", methods=["POST"])
def summarize():
    data     = request.get_json(silent=True) or {}
    filename = os.path.basename(data.get("filename","")).strip()
    print(f"\n[SUMMARIZE] '{filename}'")

    ok, err = _ensure_indexed(filename)
    if not ok:
        return jsonify({"error": err})

    chunks     = per_file_data[filename]["chunks"]
    word_count = len(" ".join(chunks).split())

    # Use chunks from the middle of the document (skip front matter)
    start = min(5, len(chunks) - 1)
    sample = chunks[start: start + 14]
    if len(sample) < 3:
        sample = chunks[:14]
    context = "\n---\n".join(sample)

    if groq_client:
        try:
            prompt = (
                "You are StudyGenie, an academic assistant.\n"
                "Read these document excerpts and provide:\n"
                "1. A clear 4-6 sentence academic summary of what this document covers\n"
                "2. Exactly 6-8 key academic topics/concepts (short noun phrases)\n"
                "3. Difficulty: Easy | Medium | Hard\n\n"
                "Ignore front matter, copyright, author names, and page numbers.\n"
                "Focus on the actual academic content.\n\n"
                "Respond ONLY with this JSON (no extra text):\n"
                '{"summary":"...","topics":["topic1","topic2",...],"difficulty":"Medium"}\n\n'
                f"Excerpts:\n{context}"
            )
            resp   = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role":"user","content":prompt}],
                temperature=0.2, max_tokens=700
            )
            raw    = resp.choices[0].message.content.strip()
            parsed = _parse_json_response(raw)
            if isinstance(parsed, dict) and "summary" in parsed:
                parsed["word_count"] = word_count
                parsed["chunks"]     = len(chunks)
                return jsonify({"status":"ok", **parsed})
        except Exception as e:
            print(f"[SUMMARIZE] Groq error: {e}")

    # Fallback extractive summary
    stop = {"this","that","with","from","have","been","they","their","will","were","also",
            "more","than","into","over","such","each","when","which","there","about","page"}
    sents = []
    for chunk in chunks[start:start+8]:
        for s in re.split(r'(?<=[.!?])\s+', chunk):
            s = s.strip()
            if 55 < len(s) < 280:
                sents.append(s)
    topics = list(dict.fromkeys(
        w for chunk in chunks[start:start+6]
        for w in re.findall(r'\b[A-Z][a-z]{3,}\b', chunk)
        if w.lower() not in stop
    ))[:8]
    return jsonify({
        "status":"ok",
        "summary": " ".join(sents[:5]) or "Document indexed and ready for Q&A.",
        "topics":  topics or ["Key Concepts","Core Topics","Definitions","Applications"],
        "difficulty":"Medium",
        "word_count":word_count,
        "chunks":len(chunks),
        "mode":"fallback"
    })


@app.route("/study-plan", methods=["POST"])
def study_plan():
    """Generate a day-by-day study schedule from the document."""
    data     = request.get_json(silent=True) or {}
    filename = os.path.basename(data.get("filename", "")).strip()
    days     = max(1, min(int(data.get("days", 5)), 30))
    hours    = data.get("hours_per_day", 2)
    print(f"\n[STUDY-PLAN] '{filename}' over {days} days")

    ok, err = _ensure_indexed(filename)
    if not ok:
        return jsonify({"error": err})

    chunks  = per_file_data[filename]["chunks"]
    sampled = _quiz_chunks(chunks, k=10)
    context = "\n\n---\n\n".join(sampled)

    # Cache
    cache_path = CACHE_DIR / (hashlib.md5(filename.encode()).hexdigest() + f"_plan_{days}.json")
    if cache_path.exists():
        try:
            with open(cache_path, encoding="utf-8") as f:
                cached = json.load(f)
            if cached.get("days"):
                print("[STUDY-PLAN] ✅ from cache")
                return jsonify(cached)
        except: pass

    if groq_client:
        prompt = (
            f"You are an expert study coach. A student has {days} days to study the document "
            f"below (about {hours} hours/day available). Create a realistic, motivating day-by-day study plan.\n\n"
            "Rules:\n"
            "- Distribute the document's key topics evenly across the days (don't cram everything on day 1)\n"
            "- Build from fundamentals → advanced as days progress\n"
            "- The final day must be revision + a full self-test\n"
            "- For each day give: a short title, 2-4 topics to cover, a one-line focus, "
            "2-3 concrete activities, and a checkpoint (what to quiz/test yourself on)\n\n"
            "Return ONLY this JSON (no extra text):\n"
            '{"days":[{"day":1,"title":"...","topics":["..."],"focus":"...",'
            '"activities":["..."],"checkpoint":"..."}],"tips":["...","..."]}\n\n'
            f"Document content:\n{context}"
        )
        for model in [GROQ_MODEL, GROQ_MODEL_FAST]:
            try:
                resp   = groq_client.chat.completions.create(
                    model=model, messages=[{"role": "user", "content": prompt}],
                    temperature=0.5, max_tokens=3000)
                parsed = _parse_json_response(resp.choices[0].message.content.strip())
                if isinstance(parsed, dict) and isinstance(parsed.get("days"), list) and parsed["days"]:
                    result = {"status": "ok", "days": parsed["days"][:days],
                              "tips": parsed.get("tips", []), "total_days": days}
                    try:
                        with open(cache_path, "w", encoding="utf-8") as f: json.dump(result, f)
                    except: pass
                    print(f"[STUDY-PLAN] ✅ via {model}")
                    return jsonify(result)
            except Exception as e:
                print(f"[STUDY-PLAN] {model} error: {e}")
                continue   # any error → try the next model before giving up to fallback

    # Fallback — split topics across days
    topics = _extract_topic_words(sampled, days * 3)
    per = max(1, len(topics) // days) if topics else 1
    plan_days = []
    for d in range(days):
        slice_ = topics[d*per:(d+1)*per] or ["Review previous material"]
        plan_days.append({
            "day": d + 1,
            "title": "Revision & Self-Test" if d == days - 1 else f"Study Session {d+1}",
            "topics": slice_,
            "focus": "Final revision and full practice quiz" if d == days - 1 else "Read, understand, and take notes",
            "activities": ["Read the relevant sections", "Summarise in your own words", "Take a practice quiz"],
            "checkpoint": "Full quiz on all topics" if d == days - 1 else f"Quiz yourself on: {', '.join(slice_[:2])}"
        })
    return jsonify({"status": "ok", "days": plan_days,
                    "tips": ["Study in focused 25-minute blocks (Pomodoro).",
                             "Review the previous day's topics before starting new ones.",
                             "Use the Chatbot to clarify anything confusing."],
                    "total_days": days, "mode": "fallback"})


@app.route("/concept-map", methods=["POST"])
def concept_map():
    """Build a concept graph (nodes + relationships) from the document."""
    data     = request.get_json(silent=True) or {}
    filename = os.path.basename(data.get("filename", "")).strip()
    print(f"\n[CONCEPT-MAP] '{filename}'")

    ok, err = _ensure_indexed(filename)
    if not ok:
        return jsonify({"error": err})

    chunks  = per_file_data[filename]["chunks"]
    sampled = _quiz_chunks(chunks, k=12)
    context = "\n\n---\n\n".join(sampled)

    # Document's most important terms (TF-IDF) — used to anchor the map on
    # what actually matters in THIS document.
    key_terms = _important_terms(filename, top_n=20)
    terms_hint = ", ".join(key_terms) if key_terms else ""

    cache_path = CACHE_DIR / (hashlib.md5(filename.encode()).hexdigest() + "_cmap2.json")
    if cache_path.exists():
        try:
            with open(cache_path, encoding="utf-8") as f:
                cached = json.load(f)
            if cached.get("nodes"):
                print("[CONCEPT-MAP] ✅ from cache")
                return jsonify(cached)
        except: pass

    if groq_client:
        prompt = (
            "Build a CONCEPT MAP of the MOST IMPORTANT topics in this document and how they "
            "relate to each other. The map should help a student see the document's key ideas "
            "at a glance.\n\n"
            "Rules:\n"
            "- Focus ONLY on the important, central topics a student must understand — "
            "ignore minor details, examples, author names, and page furniture\n"
            "- Identify 8-14 key concepts (nodes), prioritising the most significant ones\n"
            "- Put the document's main subject / overarching theme as a central node that "
            "others connect to\n"
            "- Connect related concepts with labelled relationships (edges)\n"
            "- Each edge label is a short phrase (e.g. 'is part of', 'uses', 'leads to', "
            "'type of', 'depends on')\n"
            "- Assign each node a 'group' (the broad category/theme it belongs to)\n\n"
            + (f"The document's most statistically important terms are: {terms_hint}. "
               "Use these to guide which concepts matter most (rephrase into clean concept names).\n\n"
               if terms_hint else "")
            + "Return ONLY this JSON (no extra text):\n"
            '{"nodes":[{"id":"Concept Name","group":"Category"}],'
            '"edges":[{"from":"Concept A","to":"Concept B","label":"relationship"}]}\n\n'
            f"Document content:\n{context}"
        )
        for model in [GROQ_MODEL, GROQ_MODEL_FAST]:
            try:
                resp   = groq_client.chat.completions.create(
                    model=model, messages=[{"role": "user", "content": prompt}],
                    temperature=0.4, max_tokens=2500)
                parsed = _parse_json_response(resp.choices[0].message.content.strip())
                if (isinstance(parsed, dict) and isinstance(parsed.get("nodes"), list)
                        and len(parsed["nodes"]) >= 3):
                    result = {"status": "ok",
                              "nodes": parsed["nodes"][:18],
                              "edges": parsed.get("edges", [])}
                    try:
                        with open(cache_path, "w", encoding="utf-8") as f: json.dump(result, f)
                    except: pass
                    print(f"[CONCEPT-MAP] ✅ via {model} ({len(result['nodes'])} nodes)")
                    return jsonify(result)
            except Exception as e:
                print(f"[CONCEPT-MAP] {model} error: {e}")
                continue   # any error → try the next model before giving up to fallback

    # Fallback — star graph anchored on the document's most important TF-IDF terms
    topics = key_terms[:12] or _extract_topic_words(sampled, 10)
    nodes = [{"id": "Main Topic", "group": "Core"}] + [{"id": t, "group": "Key Topic"} for t in topics]
    edges = [{"from": "Main Topic", "to": t, "label": "covers"} for t in topics]
    return jsonify({"status": "ok", "nodes": nodes, "edges": edges, "mode": "fallback"})


@app.route("/reindex-all", methods=["POST"])
def reindex_all():
    """Force re-index all documents (clears cache first)."""
    pdfs     = _list_documents()
    results  = {}
    for pdf in pdfs:
        fn = pdf.name
        # Clear cache
        _cache_path(fn).unlink(missing_ok=True)
        with _index_lock:
            per_file_data.pop(fn, None)
            index_status[fn] = "indexing"
        try:
            chunks, elapsed = _index_file(fn, pdf)
            results[fn] = {"status":"ok","chunks":chunks,"seconds":round(elapsed,2)}
        except Exception as e:
            results[fn] = {"status":"error","error":str(e)}
            with _index_lock:
                index_status[fn] = f"error:{e}"
    ok_count = sum(1 for r in results.values() if r["status"]=="ok")
    return jsonify({"indexed":ok_count,"total":len(pdfs),"results":results})


@app.route("/index-status", methods=["GET"])
def get_index_status():
    """Return per-file indexing state (ready / indexing / error:...)."""
    with _index_lock:
        snap = dict(index_status)
    ready   = [f for f,s in snap.items() if s == "ready"]
    pending = [f for f,s in snap.items() if s == "indexing"]
    errors  = {f:s for f,s in snap.items() if s.startswith("error")}
    return jsonify({
        "ready":   ready,
        "pending": pending,
        "errors":  errors,
        "total_indexed": len(ready),
    })


@app.route("/status", methods=["GET"])
def status():
    with _index_lock:
        indexed = list(per_file_data.keys())
    return jsonify({
        "status":        "running",
        "indexed_files": indexed,
        "groq_available":groq_client is not None,
        "upload_folder": str(UPLOAD_DIR),
        "file_chunks":   {f:len(d["chunks"]) for f,d in per_file_data.items()},
    })


@app.route("/groq-selftest", methods=["GET"])
def groq_selftest():
    """Make ONE real Groq call and report success or the exact error.
    'groq_available' only means the client object was created — this proves
    whether the API key actually works (invalid key / quota / dead model all
    surface here as the real exception text)."""
    if not groq_client:
        return jsonify({"ok": False,
                        "reason": "client_not_initialized",
                        "detail": "GROQ_API_KEY is empty or still the placeholder."})
    # List the models this API key is actually allowed to use.
    available = None
    try:
        available = sorted(m.id for m in groq_client.models.list().data)
    except Exception as e:
        available = f"models.list() failed: {type(e).__name__}: {e}"

    # Try a real completion with the configured model.
    try:
        r = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": "Reply with the single word: pong"}],
            max_tokens=5, temperature=0)
        return jsonify({"ok": True, "model": GROQ_MODEL,
                        "reply": r.choices[0].message.content.strip(),
                        "available_models": available})
    except Exception as e:
        return jsonify({"ok": False, "model": GROQ_MODEL,
                        "error_type": type(e).__name__, "error": str(e),
                        "available_models": available})


@app.route("/debug", methods=["GET"])
def debug():
    pdfs_on_disk = [f.name for f in _list_documents()]
    cached       = [f.name for f in CACHE_DIR.glob("*.pkl")]
    with _index_lock:
        indexed  = list(per_file_data.keys())
    return jsonify({
        "upload_folder":   str(UPLOAD_DIR),
        "pdfs_on_disk":    pdfs_on_disk,
        "cache_files":     len(cached),
        "indexed_in_memory": indexed,
        "groq_available":  groq_client is not None,
        "index_status":    dict(index_status),
    })


# ══════════════════════════════════════════════════════════════════════════════
#  STARTUP
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("  StudyGenie AI Server")
    print(f"  Upload folder : {UPLOAD_DIR}")
    print(f"  Cache folder  : {CACHE_DIR}")
    print(f"  Groq model    : {GROQ_MODEL}")
    print("=" * 60)

    # 1. Load semantic model in background (don't block server start)
    threading.Thread(target=_load_st_model, daemon=True).start()

    # 2. Index all PDFs in background (loads from disk cache = fast)
    threading.Thread(target=_auto_index_all_bg, daemon=True).start()

    print("  Server ready at http://0.0.0.0:5000")
    print("  Semantic model loading in background...")
    print("=" * 60)
    app.run(host="0.0.0.0", port=5000, threaded=True)
